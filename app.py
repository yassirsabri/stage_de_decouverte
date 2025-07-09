from flask import Flask, render_template, request, redirect, url_for, send_file
from flask_mysqldb import MySQL

import mysql.connector


from docx import Document
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
from io import BytesIO
import seaborn as sns
from scipy.stats import skew, kurtosis
import base64

app = Flask(__name__)
app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
#a changer
app.config['MYSQL_PASSWORD'] = 'your_password'
app.config['MYSQL_DB'] = 'new_schema'

def get_db_connection():
    return mysql.connector.connect(
        host=app.config['MYSQL_HOST'],
        user=app.config['MYSQL_USER'],
        password=app.config['MYSQL_PASSWORD'],
        database=app.config['MYSQL_DB']
    )

@app.route('/')
def home():
    return render_template('index.html', title="Accueil")

@app.route('/collectivites')
def collectivites():
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute('''SELECT code_poste, code_org, libelle_collectivite FROM collectivites''')
    collectivites = cur.fetchall()
    cur.close()
    conn.close()
    return render_template('data.html', title='Collectivités', column_names=['code_poste', 'code_org', 'libelle_collectivite'], data=collectivites)

@app.route('/postes_comptables')
def postes_comptables():
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute('''SELECT code_poste, libelle_poste FROM postes_comptables''')
    postes_comptables = cur.fetchall()
    cur.close()
    conn.close()
    return render_template('data.html', title='Postes Comptables', column_names=['code_poste', 'libelle_poste'], data=postes_comptables)

@app.route('/comptables')
def comptables():
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute('''SELECT code_poste, nom, prenom, email, sexe FROM comptables''')
    comptables = cur.fetchall()
    cur.close()
    conn.close()
    return render_template('data.html', title='Comptables', column_names=['code_poste', 'nom', 'prenom', 'email', 'sexe'], data=comptables)

@app.route('/nature_sinistre')
def nature_sinistre():
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("SELECT id, libelle FROM nature_sinistre")
    nature_sinistre = cur.fetchall()
    cur.close()
    conn.close()
    return render_template('data.html', title='Nature de Sinistre', column_names=['id', 'libelle'], data=nature_sinistre)

@app.route('/cours_comptes')
def cours_comptes():
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute('''SELECT id, libelle FROM cours_comptes''')
    cours_comptes = cur.fetchall()
    cur.close()
    conn.close()
    return render_template('data.html', title='Cours des Comptes', column_names=['id', 'libelle'], data=cours_comptes)

@app.route('/add_entry', methods=['GET', 'POST'])
def add_entry():
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Récupérer les postes comptables
    cursor.execute("SELECT libelle_poste FROM postes_comptables")
    postes_comptables = cursor.fetchall()

    # Récupérer les natures de sinistre
    cursor.execute("SELECT libelle FROM nature_sinistre")
    nature_sinistre = cursor.fetchall()
    
    cursor.close()
    conn.close()
    
    if request.method == 'POST':
        poste_comptable = request.form['poste_comptable']
        numero_demande = request.form['numero_demande']
        date = request.form['date']
        nature_sinistre = request.form['nature_sinistre']
        date_sinistre = request.form['date_sinistre']
        montant = request.form['montant']

        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO entries (poste_comptable, numero_demande, date, nature_sinistre, date_sinistre, montant) VALUES (%s, %s, %s, %s, %s, %s)",
                (poste_comptable, numero_demande, date, nature_sinistre, date_sinistre, montant)
            )
            conn.commit()
            cursor.close()
            conn.close()
            return redirect(url_for('confirmation'))
        except mysql.connector.Error as err:
            return f"Erreur: {err}"

    return render_template('add_entry.html', postes_comptables=postes_comptables, nature_sinistre=nature_sinistre)

@app.route('/confirmation')
def confirmation():
    return render_template('confirmation.html')

@app.route('/generate_word')
def generate_word():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM entries ORDER BY id DESC LIMIT 1")
    entry = cursor.fetchone()
    cursor.close()
    conn.close()
    
    if entry:
        doc = Document()
        doc.add_heading('Détails de l\'entrée', 0)
        
        doc.add_paragraph(f"Poste Comptable: {entry[1]}")
        doc.add_paragraph(f"Numéro de Demande: {entry[2]}")
        doc.add_paragraph(f"Date: {entry[3]}")
        doc.add_paragraph(f"Nature de Sinistre: {entry[4]}")
        doc.add_paragraph(f"Date de Sinistre: {entry[5]}")
        doc.add_paragraph(f"Montant: {entry[6]}")
        
        file_path = 'details.docx'
        doc.save(file_path)
        
        return send_file(file_path, as_attachment=True)
    
    return "Aucune entrée trouvée."

@app.route('/statistics')
def statistics():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT montant FROM entries")
    amounts = [float(row[0]) for row in cursor.fetchall()]
    cursor.close()
    conn.close()
    
    if amounts:
        amounts_array = np.array(amounts)
        
        # Calcul des statistiques avancées
        mean_amount = np.mean(amounts_array)
        median_amount = np.median(amounts_array)
        std_dev_amount = np.std(amounts_array)
        variance_amount = np.var(amounts_array)
        skewness_amount = skew(amounts_array)
        kurtosis_amount = kurtosis(amounts_array)

        # Histogram
        plt.figure(figsize=(10, 6))
        
        plt.hist(amounts_array, bins=10, color='skyblue', edgecolor='black')
        plt.title('Distribution des Montants')
        plt.xlabel('Montant')
        plt.ylabel('Fréquence')
        buffer = BytesIO()
        plt.savefig(buffer, format='png')
        buffer.seek(0)
        histogram = base64.b64encode(buffer.getvalue()).decode('utf-8')
        buffer.close()

        # Box Plot
        plt.figure(figsize=(10, 6))
        sns.boxplot(x=amounts_array)
        plt.title('Box Plot des Montants')
        buffer = BytesIO()
        plt.savefig(buffer, format='png')
        buffer.seek(0)
        boxplot = base64.b64encode(buffer.getvalue()).decode('utf-8')
        buffer.close()

        # Scatter Plot with Regression Line
        x = np.arange(len(amounts_array))
        y = amounts_array
        plt.figure(figsize=(10, 6))
        plt.scatter(x, y, color='blue')
        m, b = np.polyfit(x, y, 1)
        plt.plot(x, m*x + b, color='red')
        plt.title('Scatter Plot des Montants avec Ligne de Régression')
        plt.xlabel('Index')
        plt.ylabel('Montant')
        buffer = BytesIO()
        plt.savefig(buffer, format='png')
        buffer.seek(0)
        scatterplot = base64.b64encode(buffer.getvalue()).decode('utf-8')
        buffer.close()

        # Correlation Matrix
        df = pd.DataFrame({'amounts': amounts_array})
        corr = df.corr()
        plt.figure(figsize=(10, 6))
        sns.heatmap(corr, annot=True, cmap='coolwarm')
        plt.title('Matrice de Corrélation')
        buffer = BytesIO()
        plt.savefig(buffer, format='png')
        buffer.seek(0)
        correlation_matrix = base64.b64encode(buffer.getvalue()).decode('utf-8')
        buffer.close()
        

        return render_template('statistics.html', title='Statistiques', 
                               mean_amount=mean_amount, 
                               median_amount=median_amount, 
                               std_dev_amount=std_dev_amount, 
                               variance_amount=variance_amount,
                               skewness_amount=skewness_amount, 
                               kurtosis_amount=kurtosis_amount, 
                               histogram=histogram, 
                               boxplot=boxplot, 
                               scatterplot=scatterplot, 
                               correlation_matrix=correlation_matrix)
    
    return render_template('statistics.html', title='Statistiques', 
                           mean_amount=None, 
                           median_amount=None, 
                           std_dev_amount=None, 
                           variance_amount=None,
                           skewness_amount=None, 
                           kurtosis_amount=None, 
                           histogram=None, 
                           boxplot=None, 
                           scatterplot=None, 
                           correlation_matrix=None)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
